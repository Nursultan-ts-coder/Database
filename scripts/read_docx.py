#!/usr/bin/env python3
import argparse
import os
import sys
import shutil
import subprocess


def extract_docx_text(path: str) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: pip install python-docx"
        )

    doc = Document(path)
    lines: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Tables (tab-separated for readability)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))

    return "\n".join(lines)


def extract_doc_text_mac(path: str) -> str:
    """Extract text from legacy .doc using macOS textutil if available."""
    if shutil.which("textutil") is None:
        raise RuntimeError(
            "macOS 'textutil' not found. For .doc files, consider converting to .docx "
            "(e.g., open in Word and save as .docx) or install LibreOffice and convert."
        )
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", path],
            check=True,
            capture_output=True,
        )
        return result.stdout.decode("utf-8", errors="replace")
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Failed to extract .doc with textutil: {e}")


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Read Word file content (.docx or .doc on macOS) and output text."
    )
    parser.add_argument("path", help="Path to the Word file (.docx or .doc)")
    parser.add_argument(
        "--out",
        help="Optional output file to write the extracted text",
        default=None,
    )
    args = parser.parse_args()

    path = args.path
    if not os.path.isfile(path):
        print(f"Error: File not found: {path}", file=sys.stderr)
        return 1

    ext = os.path.splitext(path)[1].lower()

    try:
        if ext == ".docx":
            text = extract_docx_text(path)
        elif ext == ".doc":
            text = extract_doc_text_mac(path)
        else:
            print(
                f"Error: Unsupported extension '{ext}'. Use .docx or .doc (macOS)",
                file=sys.stderr,
            )
            return 1
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    if args.out:
        try:
            with open(args.out, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"Wrote extracted text to {args.out}")
        except Exception as e:
            print(f"Error writing output: {e}", file=sys.stderr)
            return 1
    else:
        # Print to stdout
        print(text)

    return 0


if __name__ == "__main__":
    sys.exit(main())
